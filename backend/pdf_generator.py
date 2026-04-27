"""
pdf_generator.py – Generates ATS-friendly PDF and DOCX resumes.
Uses Jake's Resume-style layout: clean, single-column, Times New Roman.
"""

import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    HRFlowable, Table, TableStyle,
)
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════════════
#  PDF  (ReportLab — Jake's Resume style)
# ═══════════════════════════════════════════════════════════════════════════

def _ps(fname, size, align=None, before=0, after=2, indent=0) -> ParagraphStyle:
    kw = dict(fontName=fname, fontSize=size, spaceBefore=before,
               spaceAfter=after, leftIndent=indent, textColor=colors.black)
    if align is not None:
        kw["alignment"] = align
    return ParagraphStyle(f"_{id(kw)}", **kw)


def _hr_line(thick=0.6):
    return HRFlowable(width="100%", thickness=thick,
                      color=colors.black, spaceAfter=3, spaceBefore=1)


def _two_col_table(left_para, right_para, doc_width):
    DATE_W = 1.5 * inch
    t = Table([[left_para, right_para]],
              colWidths=[doc_width - DATE_W, DATE_W])
    t.setStyle(TableStyle([
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING",   (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 0),
    ]))
    return t


def generate_ats_pdf(resume_data: dict) -> bytes:
    """Return ATS-friendly PDF bytes in Jake's Resume style."""
    buf = io.BytesIO()
    MARGIN = 0.65 * inch
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        rightMargin=MARGIN, leftMargin=MARGIN,
        topMargin=MARGIN, bottomMargin=MARGIN,
    )
    W = doc.width

    # ── Styles ────────────────────────────────────────────────────────
    s_name    = _ps("Times-Bold",   18, align=TA_CENTER, after=3)
    s_contact = _ps("Times-Roman",  9.5, align=TA_CENTER, after=4)
    s_head    = _ps("Times-Bold",   11, before=8, after=1)
    s_bold    = _ps("Times-Bold",   10, after=1)
    s_italic  = _ps("Times-Italic", 10, after=1)
    s_normal  = _ps("Times-Roman",  10, after=1)
    s_bullet  = _ps("Times-Roman",  10, after=1.5, indent=10)
    s_right   = _ps("Times-Roman",  10, align=TA_RIGHT, after=1)

    def section(title):
        return [Paragraph(title, s_head), _hr_line(0.5)]

    def row(left_text, right_text, left_style=None):
        if left_style is None:
            left_style = s_bold
        return _two_col_table(
            Paragraph(left_text, left_style),
            Paragraph(right_text, s_right),
            W,
        )

    EL = []

    # Name
    EL.append(Paragraph(resume_data.get("name", "Your Name"), s_name))

    # Contact line
    cparts = [resume_data.get(k, "") for k in ("phone", "email", "linkedin", "github", "location")]
    cparts = [p for p in cparts if p]
    if cparts:
        EL.append(Paragraph(" | ".join(cparts), s_contact))
    EL.append(_hr_line(1))

    # Summary
    if resume_data.get("summary"):
        EL += section("SUMMARY")
        EL.append(Paragraph(resume_data["summary"], s_normal))
        EL.append(Spacer(1, 4))

    # Education
    if resume_data.get("education"):
        EL += section("EDUCATION")
        for edu in resume_data["education"]:
            loc_year = " ".join(filter(None, [edu.get("location", ""), edu.get("year", "")]))
            EL.append(row(f"<b>{edu.get('institution','')}</b>", loc_year))
            deg = edu.get("degree", "")
            gpa = edu.get("gpa", "")
            EL.append(Paragraph(f"<i>{deg}" + (f", GPA: {gpa}" if gpa else "") + "</i>", s_italic))
            EL.append(Spacer(1, 4))

    # Experience
    if resume_data.get("experience"):
        EL += section("EXPERIENCE")
        for exp in resume_data["experience"]:
            EL.append(row(
                f"<b>{exp.get('title','')}</b> | {exp.get('company','')}",
                exp.get("location", ""),
            ))
            EL.append(Paragraph(f"<i>{exp.get('duration','')}</i>", s_italic))
            for b in exp.get("bullets", []):
                EL.append(Paragraph(f"• {b}", s_bullet))
            EL.append(Spacer(1, 4))

    # Projects
    if resume_data.get("projects"):
        EL += section("PROJECTS")
        for proj in resume_data["projects"]:
            tech = proj.get("technologies", "")
            label = f"<b>{proj.get('name','')}</b>" + (f" | <i>{tech}</i>" if tech else "")
            EL.append(row(label, proj.get("duration", "")))
            for b in proj.get("bullets", []):
                EL.append(Paragraph(f"• {b}", s_bullet))
            EL.append(Spacer(1, 4))

    # Skills
    if resume_data.get("skills"):
        EL += section("TECHNICAL SKILLS")
        skills = resume_data["skills"]
        if isinstance(skills, dict):
            for cat, lst in skills.items():
                sk = ", ".join(lst) if isinstance(lst, list) else str(lst)
                EL.append(Paragraph(f"<b>{cat}:</b> {sk}", s_normal))
        else:
            EL.append(Paragraph(", ".join(skills if isinstance(skills, list) else [str(skills)]), s_normal))

    # Certifications
    if resume_data.get("certifications"):
        EL += section("CERTIFICATIONS")
        certs = resume_data["certifications"]
        EL.append(Paragraph(", ".join(certs) if isinstance(certs, list) else str(certs), s_normal))

    doc.build(EL)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════
#  DOCX  (python-docx)
# ═══════════════════════════════════════════════════════════════════════════

BLACK = RGBColor(0, 0, 0)


def _no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _run(para, text, bold=False, italic=False, size=10):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = BLACK
    return r


def generate_ats_docx(resume_data: dict) -> bytes:
    """Return ATS-friendly DOCX bytes."""
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.65)

    def _center_para(text, size=10, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        _run(p, text, bold=bold, size=size)
        return p

    def _section_head(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(1)
        _run(p, title, bold=True, size=11)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "000000")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _two_col(left, right):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        tbl.columns[0].width = Inches(5.1)
        tbl.columns[1].width = Inches(1.4)
        lc, rc = tbl.cell(0, 0), tbl.cell(0, 1)
        _no_border(lc)
        _no_border(rc)
        lp, rp = lc.paragraphs[0], rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for p in (lp, rp):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
        lc.paragraphs[0].add_run(left).font.size = Pt(10)
        lc.paragraphs[0].runs[0].font.color.rgb = BLACK
        lc.paragraphs[0].runs[0].bold = True
        rc.paragraphs[0].add_run(right).font.size = Pt(10)
        rc.paragraphs[0].runs[0].font.color.rgb = BLACK

    def _italic_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        _run(p, text, italic=True)

    def _bullet_p(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        _run(p, text)

    def _normal_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        if bold_prefix:
            _run(p, f"{bold_prefix}: ", bold=True)
        _run(p, text)

    # ── Name
    _center_para(resume_data.get("name", "Your Name"), size=18, bold=True)

    # ── Contact
    cparts = [resume_data.get(k, "") for k in ("phone", "email", "linkedin", "github", "location")]
    cparts = [p for p in cparts if p]
    if cparts:
        _center_para(" | ".join(cparts), size=9)

    # ── Summary
    if resume_data.get("summary"):
        _section_head("SUMMARY")
        _normal_p(resume_data["summary"])

    # ── Education
    if resume_data.get("education"):
        _section_head("EDUCATION")
        for edu in resume_data["education"]:
            loc_year = " ".join(filter(None, [edu.get("location", ""), edu.get("year", "")]))
            _two_col(edu.get("institution", ""), loc_year)
            deg = edu.get("degree", "")
            gpa = edu.get("gpa", "")
            _italic_p(deg + (f", GPA: {gpa}" if gpa else ""))

    # ── Experience
    if resume_data.get("experience"):
        _section_head("EXPERIENCE")
        for exp in resume_data["experience"]:
            _two_col(f"{exp.get('title','')} | {exp.get('company','')}", exp.get("location", ""))
            _italic_p(exp.get("duration", ""))
            for b in exp.get("bullets", []):
                _bullet_p(b)

    # ── Projects
    if resume_data.get("projects"):
        _section_head("PROJECTS")
        for proj in resume_data["projects"]:
            tech = proj.get("technologies", "")
            label = proj.get("name", "") + (f" | {tech}" if tech else "")
            _two_col(label, proj.get("duration", ""))
            for b in proj.get("bullets", []):
                _bullet_p(b)

    # ── Skills
    if resume_data.get("skills"):
        _section_head("TECHNICAL SKILLS")
        skills = resume_data["skills"]
        if isinstance(skills, dict):
            for cat, lst in skills.items():
                sk = ", ".join(lst) if isinstance(lst, list) else str(lst)
                _normal_p(sk, bold_prefix=cat)
        else:
            _normal_p(", ".join(skills if isinstance(skills, list) else [str(skills)]))

    # ── Certifications
    if resume_data.get("certifications"):
        _section_head("CERTIFICATIONS")
        certs = resume_data["certifications"]
        _normal_p(", ".join(certs) if isinstance(certs, list) else str(certs))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
